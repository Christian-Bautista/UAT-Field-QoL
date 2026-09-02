"""Build a UAT field report from the template, shared by the desktop app and the web page."""

import copy

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SECTION_HEADING = "Signal Coverage Test"
STANDALONE_HEADINGS = ("OVERALL TEST RESULT AND SIGNOFF", "Customer Acknowledgement")
ROWS_PER_AP = 5  # AP ID, Test Device 1, SSID/Band header, 2.4 GHz, 5 GHz
ORGANISATION_LABEL = "Organisation"
VENUE_NAME_LABEL = "Venue Name"
VENUE_ADDRESS_LABEL = "Venue Address"
VENUE_POSTAL_CODE_LABEL = "Venue Postal Code"
VENUE_CATEGORY_LABEL = "Venue Category"


def find_ap_table(document):
    """Return the table under section 3.2.1 (the first table of the Signal Coverage Test section)."""
    in_section = False
    for element in document.element.body.iterchildren():
        tag = element.tag.split("}")[-1]
        if tag == "p":
            paragraph = Paragraph(element, document)
            style = paragraph.style.name
            if style == "Heading 2":
                in_section = paragraph.text.strip() == SECTION_HEADING
            elif style == "Heading 1":
                in_section = False
        elif tag == "tbl" and in_section:
            return Table(element, document)
    raise LookupError("Could not find the section 3.2.1 table in the template.")


def trim_ap_block(table):
    """Drop the template's trailing band rows so one AP block matches the full template's form."""
    for row in table._tbl.tr_lst[ROWS_PER_AP:]:
        table._tbl.remove(row)


def iter_block_items(document):
    """Yield the document body's paragraphs and tables in the order they appear."""
    for element in document.element.body.iterchildren():
        tag = element.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(element, document)
        elif tag == "tbl":
            yield Table(element, document)


def find_section_items(document, heading_text):
    """Return a Heading 1 paragraph and everything under it, up to and including its last table."""
    items = []
    started = False
    for item in iter_block_items(document):
        is_heading = isinstance(item, Paragraph) and item.style.name == "Heading 1"
        if is_heading:
            if started:
                break
            started = item.text.strip() == heading_text
        if started:
            items.append(item)
    while items and not isinstance(items[-1], Table):
        items.pop()
    if not items:
        raise LookupError("Could not find the %s section in the template." % heading_text)
    return items


def is_filler_paragraph(paragraph):
    """Report whether a paragraph is blank spacing rather than content."""
    if paragraph.text.strip():
        return False
    for tag in ("w:sectPr", "w:drawing", "w:pict", "w:object", "w:br"):
        if paragraph._p.findall(".//" + qn(tag)):
            return False
    return True


def remove_filler_before(document, paragraph):
    """Drop the blank paragraphs in front of a paragraph, so they cannot spill onto a page."""
    previous = paragraph._p.getprevious()
    while previous is not None and previous.tag == qn("w:p"):
        if not is_filler_paragraph(Paragraph(previous, document)):
            break
        earlier = previous.getprevious()
        previous.getparent().remove(previous)
        previous = earlier


def set_cant_split(row):
    """Stop a table row from being broken across two pages."""
    row_properties = row._tr.get_or_add_trPr()
    if row_properties.find(qn("w:cantSplit")) is None:
        row_properties.append(row._tr.makeelement(qn("w:cantSplit"), {}))


def set_keep_with_next(item, keep):
    """Keep a paragraph, or every paragraph of a table, with the content that follows it."""
    paragraphs = [item] if isinstance(item, Paragraph) else [
        paragraph
        for row in item.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    ]
    for paragraph in paragraphs:
        paragraph.paragraph_format.keep_with_next = keep


def keep_section_on_own_page(document, heading_text):
    """Start the section on a new page and hold its heading and tables together on it."""
    items = find_section_items(document, heading_text)
    remove_filler_before(document, items[0])
    items[0].paragraph_format.page_break_before = True
    for item in items:
        if isinstance(item, Table):
            for row in item.rows:
                set_cant_split(row)
        set_keep_with_next(item, True)
    for cell in items[-1].rows[-1].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = False


def set_cell_text(cell, value):
    """Overwrite a cell's text, keeping the formatting of its first run."""
    paragraph = cell.paragraphs[0]
    if not paragraph.runs:
        paragraph.add_run(value)
        return
    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ""


def set_labelled_value(container, label, value):
    """Replace the value cell of every table row of the container that starts with the label."""
    for table in container.tables:
        for row in table.rows:
            if row.cells[0].text.strip() == label:
                set_cell_text(row.cells[-1], value)


def apply_details(document, details):
    """Write the organisation and venue details into the first page box and the header tables."""
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.first_page_header, section.even_page_header])
    for container in containers:
        set_labelled_value(container, ORGANISATION_LABEL, details["organisation"])
        set_labelled_value(container, VENUE_NAME_LABEL, details["venue_name"])
        set_labelled_value(container, VENUE_ADDRESS_LABEL, details["venue_address"])
        set_labelled_value(container, VENUE_POSTAL_CODE_LABEL, details["venue_postal_code"])
        set_labelled_value(container, VENUE_CATEGORY_LABEL, details["venue_category"])


def build_document_steps(template, ap_count, details):
    """Build the report one stage at a time, yielding (message, fraction) before each stage."""
    yield ("Opening template...", 0.0)
    document = Document(template)

    yield ("Locating the section 3.2.1 table...", 0.1)
    table = find_ap_table(document)

    yield ("Preparing the AP block...", 0.2)
    trim_ap_block(table)
    block = [copy.deepcopy(row) for row in table._tbl.tr_lst]

    for ap_number in range(2, ap_count + 1):
        yield (
            "Adding AP %d of %d..." % (ap_number, ap_count),
            0.2 + 0.65 * (ap_number - 1) / ap_count,
        )
        for row in block:
            table._tbl.append(copy.deepcopy(row))

    yield ("Filling in the venue details...", 0.85)
    apply_details(document, details)

    yield ("Setting the page breaks...", 0.9)
    for heading in STANDALONE_HEADINGS:
        keep_section_on_own_page(document, heading)

    return document


def build_document(template, ap_count, details, report):
    """Build the report, calling report(message, fraction) as each stage starts."""
    steps = build_document_steps(template, ap_count, details)
    while True:
        try:
            message, fraction = next(steps)
        except StopIteration as finished:
            return finished.value
        report(message, fraction)
