"""Generate a UAT field report from the venue details and AP count chosen in the GUI."""

import copy
import ctypes
import os
import queue
import sys
import tempfile
import threading
import tkinter as tk
from ctypes import wintypes
from tkinter import filedialog, messagebox, ttk

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def bundled_path(name):
    """Return the path to a file shipped alongside the script or inside the packaged executable."""
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, name)


TEMPLATE = bundled_path("UAT Field Template.docx")
SECTION_HEADING = "Signal Coverage Test"
STANDALONE_HEADINGS = ("OVERALL TEST RESULT AND SIGNOFF", "Customer Acknowledgement")
ROWS_PER_AP = 5  # AP ID, Test Device 1, SSID/Band header, 2.4 GHz, 5 GHz
ORGANISATION_LABEL = "Organisation"
VENUE_NAME_LABEL = "Venue Name"
VENUE_ADDRESS_LABEL = "Venue Address"
VENUE_POSTAL_CODE_LABEL = "Venue Postal Code"
VENUE_CATEGORY_LABEL = "Venue Category"
POSTAL_CODE_LENGTH = 6

FO_MOVE = 1
FOF_SILENT = 0x0004
FOF_RENAMEONCOLLISION = 0x0008
FOF_NOCONFIRMATION = 0x0010
FOF_WANTMAPPINGHANDLE = 0x0020
FOF_NOERRORUI = 0x0400


class SHFILEOPSTRUCTW(ctypes.Structure):
    _fields_ = [
        ("hwnd", wintypes.HWND),
        ("wFunc", wintypes.UINT),
        ("pFrom", wintypes.LPCWSTR),
        ("pTo", wintypes.LPCWSTR),
        ("fFlags", ctypes.c_uint16),
        ("fAnyOperationsAborted", wintypes.BOOL),
        ("hNameMappings", ctypes.c_void_p),
        ("lpszProgressTitle", wintypes.LPCWSTR),
    ]


class SHNAMEMAPPINGW(ctypes.Structure):
    _fields_ = [
        ("pszOldPath", wintypes.LPWSTR),
        ("pszNewPath", wintypes.LPWSTR),
        ("cchOldPath", ctypes.c_int),
        ("cchNewPath", ctypes.c_int),
    ]


class HANDLETOMAPPINGS(ctypes.Structure):
    _fields_ = [
        ("uNumberOfMappings", wintypes.UINT),
        ("lpSHNameMapping", ctypes.POINTER(SHNAMEMAPPINGW)),
    ]


def move_with_os_naming(source, destination):
    """Move source onto destination, letting Windows rename it as it does for any file clash."""
    source = os.path.normpath(os.path.abspath(source))
    destination = os.path.normpath(os.path.abspath(destination))
    shell = ctypes.windll.shell32
    shell.SHFileOperationW.argtypes = [ctypes.c_void_p]
    shell.SHFileOperationW.restype = ctypes.c_int
    shell.SHFreeNameMappings.argtypes = [ctypes.c_void_p]
    shell.SHFreeNameMappings.restype = None

    operation = SHFILEOPSTRUCTW()
    operation.wFunc = FO_MOVE
    operation.pFrom = source + "\0"
    operation.pTo = destination + "\0"
    operation.fFlags = (
        FOF_RENAMEONCOLLISION
        | FOF_NOCONFIRMATION
        | FOF_SILENT
        | FOF_NOERRORUI
        | FOF_WANTMAPPINGHANDLE
    )
    result = shell.SHFileOperationW(ctypes.byref(operation))
    if result != 0:
        raise OSError("Could not save the document (Windows error %d)." % result)

    saved_path = destination
    if operation.hNameMappings:
        mappings = ctypes.cast(operation.hNameMappings, ctypes.POINTER(HANDLETOMAPPINGS)).contents
        if mappings.uNumberOfMappings:
            mapping = mappings.lpSHNameMapping[0]
            saved_path = ctypes.wstring_at(mapping.pszNewPath, mapping.cchNewPath)
        shell.SHFreeNameMappings(operation.hNameMappings)
    return saved_path


def save_document(document, output_path):
    """Save the document, leaving any file already using that name untouched."""
    folder = os.path.dirname(os.path.abspath(output_path))
    handle, temporary_path = tempfile.mkstemp(suffix=".docx", dir=folder)
    os.close(handle)
    try:
        document.save(temporary_path)
        return move_with_os_naming(temporary_path, output_path)
    finally:
        if os.path.exists(temporary_path):
            os.remove(temporary_path)


def find_ap_table(document):
    """Return the table under section 3.2.1 (the first table of the Signal Coverage Test section)."""
    in_section = False
    for element in document.element.body.iterchildren():
        tag = element.tag.split("}")[-1]
        if tag == "p":
            paragraph = Paragraph(element, document)
            style = paragraph.style.name
            if style == "Heading 2":
                in_section = paragraph.text.strip() == SECTION_HEADING
            elif style == "Heading 1":
                in_section = False
        elif tag == "tbl" and in_section:
            return Table(element, document)
    raise LookupError("Could not find the section 3.2.1 table in the template.")


def trim_ap_block(table):
    """Drop the template's trailing band rows so one AP block matches the full template's form."""
    for row in table._tbl.tr_lst[ROWS_PER_AP:]:
        table._tbl.remove(row)


def iter_block_items(document):
    """Yield the document body's paragraphs and tables in the order they appear."""
    for element in document.element.body.iterchildren():
        tag = element.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(element, document)
        elif tag == "tbl":
            yield Table(element, document)


def find_section_items(document, heading_text):
    """Return a Heading 1 paragraph and everything under it, up to and including its last table."""
    items = []
    started = False
    for item in iter_block_items(document):
        is_heading = isinstance(item, Paragraph) and item.style.name == "Heading 1"
        if is_heading:
            if started:
                break
            started = item.text.strip() == heading_text
        if started:
            items.append(item)
    while items and not isinstance(items[-1], Table):
        items.pop()
    if not items:
        raise LookupError("Could not find the %s section in the template." % heading_text)
    return items


def is_filler_paragraph(paragraph):
    """Report whether a paragraph is blank spacing rather than content."""
    if paragraph.text.strip():
        return False
    for tag in ("w:sectPr", "w:drawing", "w:pict", "w:object", "w:br"):
        if paragraph._p.findall(".//" + qn(tag)):
            return False
    return True


def remove_filler_before(document, paragraph):
    """Drop the blank paragraphs in front of a paragraph, so they cannot spill onto a page."""
    previous = paragraph._p.getprevious()
    while previous is not None and previous.tag == qn("w:p"):
        if not is_filler_paragraph(Paragraph(previous, document)):
            break
        earlier = previous.getprevious()
        previous.getparent().remove(previous)
        previous = earlier


def set_cant_split(row):
    """Stop a table row from being broken across two pages."""
    row_properties = row._tr.get_or_add_trPr()
    if row_properties.find(qn("w:cantSplit")) is None:
        row_properties.append(row._tr.makeelement(qn("w:cantSplit"), {}))


def set_keep_with_next(item, keep):
    """Keep a paragraph, or every paragraph of a table, with the content that follows it."""
    paragraphs = [item] if isinstance(item, Paragraph) else [
        paragraph
        for row in item.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    ]
    for paragraph in paragraphs:
        paragraph.paragraph_format.keep_with_next = keep


def keep_section_on_own_page(document, heading_text):
    """Start the section on a new page and hold its heading and tables together on it."""
    items = find_section_items(document, heading_text)
    remove_filler_before(document, items[0])
    items[0].paragraph_format.page_break_before = True
    for item in items:
        if isinstance(item, Table):
            for row in item.rows:
                set_cant_split(row)
        set_keep_with_next(item, True)
    for cell in items[-1].rows[-1].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = False


def set_cell_text(cell, value):
    """Overwrite a cell's text, keeping the formatting of its first run."""
    paragraph = cell.paragraphs[0]
    if not paragraph.runs:
        paragraph.add_run(value)
        return
    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ""


def set_labelled_value(container, label, value):
    """Replace the value cell of every table row of the container that starts with the label."""
    for table in container.tables:
        for row in table.rows:
            if row.cells[0].text.strip() == label:
                set_cell_text(row.cells[-1], value)


def apply_details(document, details):
    """Write the organisation and venue details into the first page box and the header tables."""
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.first_page_header, section.even_page_header])
    for container in containers:
        set_labelled_value(container, ORGANISATION_LABEL, details["organisation"])
        set_labelled_value(container, VENUE_NAME_LABEL, details["venue_name"])
        set_labelled_value(container, VENUE_ADDRESS_LABEL, details["venue_address"])
        set_labelled_value(container, VENUE_POSTAL_CODE_LABEL, details["venue_postal_code"])
        set_labelled_value(container, VENUE_CATEGORY_LABEL, details["venue_category"])


def generate_document(ap_count, details, output_path, report):
    """Write a report holding ap_count AP blocks, calling report(message, fraction) as it goes."""
    report("Opening template...", 0.0)
    document = Document(TEMPLATE)

    report("Locating the section 3.2.1 table...", 0.1)
    table = find_ap_table(document)

    report("Preparing the AP block...", 0.2)
    trim_ap_block(table)
    block = [copy.deepcopy(row) for row in table._tbl.tr_lst]

    for ap_number in range(2, ap_count + 1):
        report(
            "Adding AP %d of %d..." % (ap_number, ap_count),
            0.2 + 0.65 * (ap_number - 1) / ap_count,
        )
        for row in block:
            table._tbl.append(copy.deepcopy(row))

    report("Filling in the venue details...", 0.85)
    apply_details(document, details)

    report("Setting the page breaks...", 0.9)
    for heading in STANDALONE_HEADINGS:
        keep_section_on_own_page(document, heading)

    report("Saving document...", 0.95)
    saved_path = save_document(document, output_path)
    report("Done.", 1.0)
    return saved_path


class APQuantityPreset:
    """The window that collects the report details and reports generation progress."""

    def __init__(self, root):
        self.root = root
        self.root.title("AP Quantity Preset")
        self.save_folder = os.getcwd()
        self.updates = queue.Queue()
        self.postal_code_filled = False
        self.entries = []

        frame = ttk.Frame(root, padding=12)
        frame.grid(sticky="nsew")

        self.location_button = ttk.Button(frame, text="File Location", command=self.choose_folder)
        self.location_button.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 12))

        self.ap_count = self.add_field(frame, 1, "How many APs to be tested: ")
        self.name = self.add_field(frame, 2, "Name of document: ")
        self.organisation = self.add_field(frame, 3, "Organisation: ")
        self.venue_name = self.add_field(frame, 4, "Venue Name: ")
        self.venue_address = self.add_field(frame, 5, "Venue Address: ")
        self.venue_postal_code = self.add_field(frame, 6, "Venue Postal Code: ")
        self.venue_category = self.add_field(frame, 7, "Venue Category: ")
        self.venue_address.trace_add("write", self.on_venue_address_changed)

        actions = ttk.Frame(frame)
        actions.grid(row=8, column=0, columnspan=2, sticky="w", pady=(12, 0))

        self.generate_button = ttk.Button(actions, text="Generate Document", command=self.generate)
        self.generate_button.grid(row=0, column=0, sticky="w")

        self.progress = ttk.Progressbar(actions, length=160, maximum=100)
        self.progress.grid(row=0, column=1, padx=(12, 8))

        self.status = ttk.Label(actions, text="", width=26, anchor="w")
        self.status.grid(row=0, column=2, sticky="w")

        outline = tk.Frame(actions, background="red")
        outline.grid(row=1, column=0, sticky="w", pady=(8, 0))
        self.reset_button = tk.Button(
            outline,
            text="Reset Fields",
            command=self.reset_fields,
            foreground="red",
            activeforeground="red",
            disabledforeground="#e08080",
            relief="flat",
            borderwidth=0,
            highlightthickness=0,
            padx=8,
            pady=2,
        )
        self.reset_button.pack(padx=1, pady=1)

    def add_field(self, frame, row, label):
        """Add a prompt and its entry on the given row, returning the entry's variable."""
        pady = (8, 0) if row > 1 else (0, 0)
        ttk.Label(frame, text=label).grid(row=row, column=0, sticky="w", pady=pady)
        variable = tk.StringVar()
        entry = ttk.Entry(frame, width=28, textvariable=variable)
        entry.grid(row=row, column=1, sticky="w", pady=pady)
        self.entries.append(entry)
        return variable

    def on_venue_address_changed(self, *_):
        """Fill the postal code once the address ends in a space followed by six digits."""
        if self.postal_code_filled:
            return
        tail = self.venue_address.get()[-(POSTAL_CODE_LENGTH + 1):]
        if len(tail) == POSTAL_CODE_LENGTH + 1 and tail[0] == " " and tail[1:].isdigit():
            self.venue_postal_code.set(tail[1:])
            self.postal_code_filled = True

    def reset_fields(self):
        """Clear every field."""
        self.postal_code_filled = False
        for variable in (
            self.ap_count,
            self.name,
            self.organisation,
            self.venue_name,
            self.venue_address,
            self.venue_postal_code,
            self.venue_category,
        ):
            variable.set("")

    def choose_folder(self):
        self.postal_code_filled = False
        folder = filedialog.askdirectory(initialdir=self.save_folder, title="File Location")
        if folder:
            self.save_folder = folder

    def set_inputs_enabled(self, enabled):
        state = "normal" if enabled else "disabled"
        for entry in self.entries:
            entry.configure(state=state)
        self.generate_button.configure(state=state)
        self.location_button.configure(state=state)
        self.reset_button.configure(state=state)

    def generate(self):
        self.postal_code_filled = False
        try:
            ap_count = int(self.ap_count.get())
        except ValueError:
            messagebox.showerror("AP Quantity Preset", "Please enter a whole number of APs.")
            return
        if ap_count < 1:
            messagebox.showerror("AP Quantity Preset", "The number of APs must be at least 1.")
            return

        name = self.name.get().strip()
        if not name:
            messagebox.showerror("AP Quantity Preset", "Please enter a name for the document.")
            return
        if not name.lower().endswith(".docx"):
            name += ".docx"
        output_path = os.path.join(self.save_folder, name)

        details = {
            "organisation": self.organisation.get().strip(),
            "venue_name": self.venue_name.get().strip(),
            "venue_address": self.venue_address.get().strip(),
            "venue_postal_code": self.venue_postal_code.get().strip(),
            "venue_category": self.venue_category.get().strip(),
        }

        self.set_inputs_enabled(False)
        self.progress.configure(value=0)
        self.status.configure(text="Starting...")

        worker = threading.Thread(
            target=self.run_generation, args=(ap_count, details, output_path), daemon=True
        )
        worker.start()
        self.root.after(100, self.drain_updates)

    def run_generation(self, ap_count, details, output_path):
        try:
            saved_path = generate_document(
                ap_count,
                details,
                output_path,
                lambda message, fraction: self.updates.put(("progress", message, fraction)),
            )
        except Exception as error:  # reported in the window rather than the console
            self.updates.put(("error", str(error), 0.0))
        else:
            self.updates.put(("finished", saved_path, 1.0))

    def drain_updates(self):
        finished = False
        while True:
            try:
                kind, message, fraction = self.updates.get_nowait()
            except queue.Empty:
                break
            self.progress.configure(value=fraction * 100)
            if kind == "progress":
                self.status.configure(text=message)
            elif kind == "error":
                self.status.configure(text="Failed.")
                messagebox.showerror("AP Quantity Preset", message)
                finished = True
            else:
                self.status.configure(text="Saved %s" % os.path.basename(message))
                finished = True
        if finished:
            self.set_inputs_enabled(True)
        else:
            self.root.after(100, self.drain_updates)


def main():
    root = tk.Tk()
    APQuantityPreset(root)
    root.mainloop()


if __name__ == "__main__":
    main()
